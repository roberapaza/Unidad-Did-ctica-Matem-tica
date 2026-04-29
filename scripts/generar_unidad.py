from __future__ import annotations

import json
import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
HEADER_FILL = "D9E2F3"
HEADER_DARK = "8EAADB"
BORDER_COLOR = "000000"


def load_data(path: Path) -> dict[str, Any]:
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def set_run_font(run, *, size=11, bold=False, italic=False, underline=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_font(paragraph, size=11, bold=False, color=None):
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold, color=color)


def add_paragraph(doc: Document, text: str = "", *, align=None, size=11, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, italic=italic, color=color)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.0
    return p


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, *, bold=False, size=11, color=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    for i, line in enumerate(str(text).split("\n")):
        if i:
            p.add_run().add_break()
        r = p.add_run(line)
        set_run_font(r, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color=BORDER_COLOR, sz="6"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), sz)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def shade_header_row(table, fill=HEADER_FILL):
    for cell in table.rows[0].cells:
        set_cell_shading(cell, fill)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True


def add_table(doc: Document, headers: list[str], rows: list[list[str]], *, header_fill=HEADER_FILL):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    shade_header_row(table, header_fill)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_section_heading(doc: Document, text: str):
    p = add_paragraph(doc, text, size=11, bold=True)
    p.paragraph_format.space_before = Pt(6)
    return p


def add_header_images(doc: Document):
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    table = header.add_table(rows=1, cols=2, width=Cm(25.0))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    left = table.cell(0, 0)
    right = table.cell(0, 1)
    left_p = left.paragraphs[0]
    right_p = right.paragraphs[0]
    left_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    inst = ROOT / "recursos" / "logo_institucional.png"
    khan = ROOT / "recursos" / "logo_khan_academy.jpeg"
    if inst.exists():
        left_p.add_run().add_picture(str(inst), width=Cm(8.5))
    else:
        r = left_p.add_run("Institución Educativa N.° 0148")
        set_run_font(r, size=10, bold=True)
    if khan.exists():
        right_p.add_run().add_picture(str(khan), width=Cm(2.0))
    else:
        r = right_p.add_run("Khan Academy")
        set_run_font(r, size=10, bold=True)
    set_table_borders(table, color="FFFFFF", sz="0")


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(2.83)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(0.4)
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    styles["Normal"].font.size = Pt(11)
    add_header_images(doc)
    return doc


def add_title_block(doc: Document, data: dict[str, Any]):
    p = add_paragraph(doc, f"UNIDAD DE APRENDIZAJE N° {data['numero_unidad']}", align=WD_ALIGN_PARAGRAPH.CENTER, size=20, bold=True, color="000000")
    for run in p.runs:
        run.underline = True
    add_paragraph(doc, f"“{data['titulo_unidad']}”", align=WD_ALIGN_PARAGRAPH.CENTER, size=20, bold=True, color="0000FF")


def add_datos_informativos(doc: Document, data: dict[str, Any]):
    items = list(data["datos_informativos"].items())
    table = doc.add_table(rows=0, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)
    for i in range(0, len(items), 2):
        row = table.add_row().cells
        k1, v1 = items[i]
        set_cell_text(row[0], k1.upper(), bold=True)
        set_cell_shading(row[0], HEADER_FILL)
        set_cell_text(row[1], v1)
        if i + 1 < len(items):
            k2, v2 = items[i + 1]
            set_cell_text(row[2], k2.upper(), bold=True)
            set_cell_shading(row[2], HEADER_FILL)
            set_cell_text(row[3], v2)
        else:
            set_cell_text(row[2], "")
            set_cell_text(row[3], "")
    doc.add_paragraph()


def add_situacion(doc: Document, data: dict[str, Any]):
    add_section_heading(doc, "I. SITUACIÓN SIGNIFICATIVA:")
    table = doc.add_table(rows=2, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    set_cell_text(table.cell(0, 0), "SITUACIÓN SIGNIFICATIVA", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(table.cell(0, 0), HEADER_FILL)
    texto = data["situacion_significativa"] + " " + " ".join(data.get("preguntas_retadoras", []))
    set_cell_text(table.cell(1, 0), texto)
    doc.add_paragraph()

    contexto = data["contexto"]
    headers = ["PROBLEMÁTICA DE CONTEXTO", "NECESIDADES E INTERESES DE LOS ESTUDIANTES", "COMPETENCIAS SELECCIONADAS"]
    rows = [[contexto["Problemática de contexto"], contexto["Necesidades e intereses de los estudiantes"], contexto["Competencias seleccionadas"]]]
    add_table(doc, headers, rows)
    add_table(doc, ["PRODUCTO DE LA UNIDAD"], [[contexto["Producto de la unidad"]]])


def add_enfoques(doc: Document, data: dict[str, Any]):
    add_section_heading(doc, "II. ENFOQUES TRANSVERSALES")
    rows = [[e["enfoque"], e["acciones"]] for e in data["enfoques_transversales"]]
    add_table(doc, ["Enfoques transversales", "Acciones o actitudes"], rows)


def add_comp_transversales(doc: Document, data: dict[str, Any]):
    add_section_heading(doc, "III. COMPETENCIAS TRANSVERSALES")
    table = add_table(doc, ["Com. transversales", "Capacidades", "Criterios de evaluación"], [])
    for item in data["competencias_transversales"]:
        row = table.add_row().cells
        set_cell_text(row[0], item["competencia"])
        set_cell_text(row[1], item["capacidad"])
        # Rich text for Khan Academy phrase
        row[2].text = ""
        p = row[2].paragraphs[0]
        text = item["criterio"]
        key = "así como reforzar sus aprendizajes en la plataforma Khan Academy"
        if key in text:
            before = text.split(key)[0]
            after = text.split(key)[1]
            r1 = p.add_run(before)
            set_run_font(r1)
            r2 = p.add_run(key)
            set_run_font(r2, bold=True, color="EE0000")
            r3 = p.add_run(after)
            set_run_font(r3)
        else:
            r = p.add_run(text)
            set_run_font(r)
    set_table_borders(table)
    doc.add_paragraph()


def add_matriz(doc: Document, data: dict[str, Any]):
    add_section_heading(doc, "IV. MATRIZ DE PLANIFICACIÓN: ESTÁNDAR DE APRENDIZAJE / COMPETENCIA / CAPACIDADES / CRITERIOS DE EVALUACIÓN / PRODUCTOS E INSTRUMENTOS")
    for comp in data["competencias_area"]:
        add_table(doc, ["ÁREA", "COMPETENCIA"], [[comp["area"], comp["competencia"]]], header_fill=HEADER_DARK)
        add_table(doc, ["ESTÁNDAR - CICLO VI"], [[comp["estandar"]]])
        rows = [[f["capacidad"], f["criterio"], f["evidencia"], f["instrumento"]] for f in comp["filas"]]
        add_table(doc, ["CAPACIDADES", "CRITERIOS DE EVALUACIÓN", "EVIDENCIAS DE PROCESO", "INSTRUMENTOS DE EVALUACIÓN"], rows)


def add_sesiones(doc: Document, data: dict[str, Any]):
    add_section_heading(doc, "V. SECUENCIA DE SESIONES")
    sesiones = data["sesiones"]
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)
    for i in range(0, len(sesiones), 2):
        row = table.add_row().cells
        for j in range(2):
            cell = row[j]
            if i + j < len(sesiones):
                s = sesiones[i + j]
                set_cell_shading(cell, HEADER_DARK)
                cell.text = ""
                p = cell.paragraphs[0]
                r = p.add_run(f"SESIÓN {s['numero']}\n")
                set_run_font(r, bold=True)
                r = p.add_run(f"Título: “{s['titulo']}”\n\n")
                set_run_font(r, bold=True)
                r = p.add_run(f"Propósito: {s['proposito']}")
                set_run_font(r)
            else:
                set_cell_text(cell, "")
    doc.add_paragraph()


def add_final_sections(doc: Document, data: dict[str, Any]):
    add_section_heading(doc, "VI. PRODUCTO INTEGRADO DE LA UNIDAD:")
    add_paragraph(doc, data["producto_integrado"])

    add_section_heading(doc, "VII. INSTRUMENTOS DE EVALUACIÓN:")
    add_paragraph(doc, ", ".join(data["instrumentos"]))

    add_section_heading(doc, "VIII. MATERIALES Y RECURSOS:")
    rows = []
    max_len = max(len(data["materiales"]["Para el docente"]), len(data["materiales"]["Para el estudiante"]))
    for i in range(max_len):
        docente = data["materiales"]["Para el docente"][i] if i < len(data["materiales"]["Para el docente"]) else ""
        estudiante = data["materiales"]["Para el estudiante"][i] if i < len(data["materiales"]["Para el estudiante"]) else ""
        rows.append([docente, estudiante])
    add_table(doc, ["PARA EL DOCENTE", "PARA EL ESTUDIANTE"], rows)

    add_section_heading(doc, "DOCENTES RESPONSABLES:")
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    for i, firma in enumerate(data["firmas"]):
        cell = table.cell(0, i)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if firma.get("firma"):
            img_path = ROOT / firma["firma"]
            if img_path.exists():
                try:
                    p.add_run().add_picture(str(img_path), width=Cm(3.0))
                    p.add_run().add_break()
                except Exception:
                    pass
        r = p.add_run("_______________________________\n")
        set_run_font(r)
        r = p.add_run(firma["nombre"] + "\n")
        set_run_font(r, bold=True)
        r = p.add_run(firma["cargo"])
        set_run_font(r)


def generate(data_path: Path) -> Path:
    data = load_data(data_path)
    doc = setup_document()
    add_title_block(doc, data)
    add_datos_informativos(doc, data)
    add_situacion(doc, data)
    add_enfoques(doc, data)
    add_comp_transversales(doc, data)
    add_matriz(doc, data)
    add_sesiones(doc, data)
    add_final_sections(doc, data)

    out_path = ROOT / data.get("salida", f"salida/UNIDAD_{data['numero_unidad']}.docx")
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Uso: python scripts/generar_unidad.py datos/unidad_2.json", file=sys.stderr)
        raise SystemExit(2)
    data_path = Path(sys.argv[1])
    if not data_path.is_absolute():
        data_path = ROOT / data_path
    out = generate(data_path)
    print(f"Documento generado: {out}")
